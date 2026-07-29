"""生成 P5 测试样本：4 种格式各一份样本（含表格 / 标题 / 多 sheet 等）。

跑：
    python projects/p5-multi-format-loader/scripts/make_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT = Path(__file__).resolve().parents[1] / "fixtures"
OUT.mkdir(parents=True, exist_ok=True)


def make_pdf() -> Path:
    out = OUT / "sample.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=letter)
    styles = getSampleStyleSheet()

    rev_tbl = Table(
        data=[
            ["Region", "Revenue (M USD)", "Growth"],
            ["North America", "320", "12%"],
            ["Europe", "180", "9%"],
            ["APAC", "240", "32%"],
            ["LATAM", "95", "15%"],
        ],
        colWidths=[150, 150, 100],
    )
    rev_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )

    # ── 真实图片：APAC 营收图（含 OCR 关键词 "APAC Revenue 240M"）──────────
    here = Path(__file__).resolve().parent
    assets = here.parent / "assets"
    apac_img = Image(str(assets / "chart_apac.png"), width=440, height=248)
    growth_img = Image(str(assets / "chart_growth.png"), width=440, height=220)

    flow = [
        Paragraph("Quarterly Performance Report", styles["Title"]),
        Spacer(1, 12),
        Paragraph("Q3 2024 Business Review", styles["Heading1"]),
        Paragraph(
            "This report summarizes the key metrics for the third quarter of 2024. "
            "Revenue grew 18% YoY, driven by strong performance in the APAC region.",
            styles["BodyText"],
        ),
        Spacer(1, 12),
        Paragraph("Revenue by Region", styles["Heading2"]),
        rev_tbl,
        Spacer(1, 18),
        Paragraph("Figure 1 - APAC Revenue Breakdown", styles["Heading3"]),
        apac_img,
        Spacer(1, 24),
        Paragraph("Conclusion", styles["Heading2"]),
        Paragraph(
            "APAC remains the highest-growth region. Recommend doubling APAC investment in Q4.",
            styles["BodyText"],
        ),
        PageBreak(),
        Paragraph("Regional Growth Rate", styles["Heading1"]),
        Paragraph(
            "Figure 2 below shows the growth rate per region.",
            styles["BodyText"],
        ),
        growth_img,
        Spacer(1, 12),
        Paragraph("Methodology", styles["Heading2"]),
        Paragraph("Revenue figures use constant currency. Growth is YoY.", styles["BodyText"]),
    ]
    doc.build(flow)
    return out


def make_docx() -> Path:
    out = OUT / "sample.docx"
    doc = DocxDocument()
    doc.add_heading("Employee Onboarding Checklist", 0)
    doc.add_paragraph(
        "This document outlines the first 30 days for new hires in the engineering team."
    )
    doc.add_heading("Week 1 Tasks", 1)
    doc.add_paragraph("Complete HR onboarding and IT setup within 2 business days.")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light List Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Day"
    hdr[1].text = "Task"
    hdr[2].text = "Owner"
    for row in [
        ["Day 1", "Receive laptop and accounts", "IT"],
        ["Day 2", "Meet team and manager", "Manager"],
        ["Day 3", "Codebase walkthrough", "Tech Lead"],
        ["Day 5", "First PR (good first issue)", "Self"],
    ]:
        c = tbl.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v
    doc.add_heading("Week 2 Goals", 1)
    doc.add_paragraph("Ship a small feature to production.")
    doc.add_heading("Week 4 Goals", 1)
    doc.add_paragraph("Own a component end-to-end with mentorship.")
    doc.save(str(out))
    return out


def make_xlsx() -> Path:
    out = OUT / "sample.xlsx"
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Headcount"
    ws1.append(["Department", "Q1", "Q2", "Q3", "Q4"])
    ws1.append(["Engineering", 42, 45, 48, 51])
    ws1.append(["Sales", 18, 20, 22, 24])
    ws1.append(["Operations", 12, 12, 13, 14])
    ws1.append(["Marketing", 8, 9, 10, 11])

    ws2 = wb.create_sheet("Budget")
    ws2.append(["Category", "Allocated (k)", "Spent (k)", "Variance"])
    ws2.append(["Salaries", 4200, 4100, 100])
    ws2.append(["Infrastructure", 800, 920, -120])
    ws2.append(["Marketing", 600, 540, 60])

    wb.save(str(out))
    return out


def make_html() -> Path:
    out = OUT / "sample.html"
    out.write_text(
        """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Product Catalog</title>
  <style>body { font-family: sans-serif; }</style>
</head>
<body>
  <nav><a href="/">Home</a> · <a href="/about">About</a></nav>
  <article>
    <h1>Product Catalog</h1>
    <p>This catalog lists our flagship products available in 2026.</p>

    <h2>Flagship Products</h2>
    <p>Our flagship product line combines performance with reliability.</p>

    <table>
      <thead>
        <tr><th>SKU</th><th>Name</th><th>Price</th><th>Stock</th></tr>
      </thead>
      <tbody>
        <tr><td>A-001</td><td>Pro Laptop</td><td>$1499</td><td>120</td></tr>
        <tr><td>A-002</td><td>Ultra Tablet</td><td>$899</td><td>240</td></tr>
        <tr><td>B-101</td><td>Smart Hub</td><td>$249</td><td>510</td></tr>
      </tbody>
    </table>

    <h2>Availability</h2>
    <p>Ships within 2 business days in North America.</p>
    <img src="hero.jpg" alt="Pro Laptop hero image">

    <footer>
      <p>Contact: sales@example.com</p>
    </footer>
  </article>
</body>
</html>
""",
        encoding="utf-8",
    )
    return out


if __name__ == "__main__":
    paths = [make_pdf(), make_docx(), make_xlsx(), make_html()]
    for p in paths:
        print(f"  ✓ {p.name:15s} {p.stat().st_size:6d} bytes")
    print(f"\nfixture 目录: {OUT}")
