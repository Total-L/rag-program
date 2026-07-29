"""Word (.docx) loader —— 段落 + 表格。

策略：
- python-docx 抽段落（按 document body 顺序），heading 单独打 HEADING 标
- 表格按出现位置内联抽，转 markdown
- 表格内图片：docx 嵌入图存在 document.part.related_parts 里，挨个遍历即可（不 OCR）
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument

from ..models import Block, BlockType, Document, ExtractionError, SourceFormat, _id_for

log = logging.getLogger(__name__)


def _table_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    n_cols = max((len(r) for r in rows), default=0)
    if n_cols == 0:
        return ""
    norm = [(r + [""] * n_cols)[:n_cols] for r in rows]
    lines = ["| " + " | ".join(norm[0]) + " |", "|" + "|".join(["---"] * n_cols) + "|"]
    for r in norm[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def load(path: Path) -> Document:
    if path.suffix.lower() not in (".docx", ".doc"):
        raise ExtractionError(f"DOCX loader got {path.suffix}")
    path = Path(path).resolve()
    blocks: list[Block] = []

    doc = DocxDocument(str(path))
    page_idx = 1  # docx 没真分页，按段落累计字数估
    char_count = 0

    # 1. 段落 + 表格按 body 顺序抽
    from docx.oxml.ns import qn

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            # 找对应段落对象
            text_parts: list[str] = []
            for r in child.iter(qn("w:t")):
                text_parts.append(r.text or "")
            text = "".join(text_parts).strip()
            if not text:
                continue
            # 判断 heading（python-docx 1.x 没有 WD_PARAGRAPH_STYLE，用 style.name 字符串判断）
            try:
                para_obj = next((p for p in doc.paragraphs if p._element is child), None)
                style = para_obj.style.name if para_obj else ""
            except Exception:
                style = ""
            btype = BlockType.HEADING if "Heading" in style else BlockType.TEXT
            char_count += len(text)
            blocks.append(
                Block(
                    block_id=_id_for(path, page_idx, btype, len(blocks)),
                    source_path=str(path),
                    source_format=SourceFormat.DOCX,
                    block_type=btype,
                    text=text,
                    page_or_sheet=page_idx,
                    metadata={
                        "heading_level": style.replace("Heading ", "") if "Heading" in style else ""
                    },
                )
            )
            if char_count > 1500:  # 估算换页
                page_idx += 1
                char_count = 0
        elif tag == "tbl":
            tbl_obj = next((t for t in doc.tables if t._element is child), None)
            if tbl_obj is None:
                continue
            rows: list[list[str]] = []
            for r in tbl_obj.rows:
                rows.append([c.text.strip() for c in r.cells])
            if not rows:
                continue
            md = _table_to_markdown(rows)
            headers = rows[0]
            blocks.append(
                Block(
                    block_id=_id_for(path, page_idx, BlockType.TABLE, len(blocks)),
                    source_path=str(path),
                    source_format=SourceFormat.DOCX,
                    block_type=BlockType.TABLE,
                    text=md,
                    page_or_sheet=page_idx,
                    headers=tuple(headers),
                    rows=tuple(tuple(r) for r in rows[1:]),
                )
            )

    # 2. 文档级图片（part 关系里的所有 image）
    try:
        img_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_count += 1
                blocks.append(
                    Block(
                        block_id=_id_for(path, "embedded", BlockType.IMAGE, len(blocks)),
                        source_path=str(path),
                        source_format=SourceFormat.DOCX,
                        block_type=BlockType.IMAGE,
                        text="",
                        page_or_sheet="embedded",
                        metadata={"rel_id": rel.rId, "target": rel.target_ref},
                    )
                )
    except Exception as e:  # noqa: BLE001
        log.warning("docx image scan failed: %s", e)

    return Document(
        source_path=str(path),
        source_format=SourceFormat.DOCX,
        blocks=tuple(blocks),
    )
